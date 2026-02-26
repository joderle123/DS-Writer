"""
Cover page and table of contents for the Praxishandbuch.
"""
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def add_cover_page(doc):
    """Add a professional cover page."""
    # Add empty lines for spacing
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        # Prevent page break before
        pPr = p._element.get_or_add_pPr()
        pageBreak = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
        pPr.append(pageBreak)

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pageBreak = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
    pPr.append(pageBreak)
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="003366"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pageBreak2 = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
    pPr.append(pageBreak2)
    run = p.add_run("PRAXISHANDBUCH")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pageBreak3 = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
    pPr.append(pageBreak3)
    run = p.add_run("Arbeit mit Jugendlichen mit\nsozio-emotionalen Schwierigkeiten")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0, 76, 153)
    run.font.name = 'Calibri'

    # Horizontal line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = p._element.get_or_add_pPr()
    pageBreak4 = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
    pPr.append(pageBreak4)
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:top w:val="single" w:sz="12" w:space="1" w:color="003366"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # Spacer
    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CDSE - Centre de développement scolaire et d'éducation")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 102, 153)
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Luxemburg")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 102, 153)

    # Spacer
    for _ in range(3):
        doc.add_paragraph()

    # Description
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Ein umfassendes Referenzdokument für Psychologen im pädagogisch-psychologischen Dienst\n"
        "Beratung · Prävention · Förderung · Krisenintervention"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)
    run.font.italic = True

    # Spacer
    for _ in range(2):
        doc.add_paragraph()

    # Footer info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Zielgruppe: Jugendliche 10-18 Jahre")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Version 1.0 - 2026")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Page break after cover
    doc.add_page_break()


def add_table_of_contents(doc):
    """Add table of contents page."""
    p = doc.add_paragraph()
    run = p.add_run("INHALTSVERZEICHNIS")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Suppress automatic page break
    pPr = p._element.get_or_add_pPr()
    pageBreak = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
    pPr.append(pageBreak)

    doc.add_paragraph()

    toc_entries = [
        ("KAPITEL 1", "GRUNDLAGEN DER ADOLESZENZ", [
            "1.1 Neurobiologische Entwicklung des jugendlichen Gehirns",
            "1.2 Psychosoziale Entwicklungsaufgaben",
            "1.3 Bindungstheorie im Jugendalter",
            "1.4 Sozio-emotionale Entwicklung",
        ]),
        ("KAPITEL 2", "STÖRUNGSBILDER UND AUFFÄLLIGKEITEN", [
            "2.1 Angststörungen bei Jugendlichen",
            "2.2 Depression bei Jugendlichen",
            "2.3 Selbstverletzendes Verhalten und Suizidalität",
            "2.4 Aggression und oppositionelles Verhalten",
            "2.5 ADHS im Jugendalter",
            "2.6 Autismus-Spektrum",
            "2.7 Trauma und PTBS",
            "2.8 Essstörungen",
            "2.9 Substanzkonsum und Verhaltenssüchte",
            "2.10 Mobbing und Cybermobbing",
            "2.11 Schulverweigerung und Schulabsentismus",
            "2.12 Gender, Sexualität und Identität",
            "2.13 Migration, Flucht und kulturelle Identität",
            "2.14 Familiäre Belastungen",
        ]),
        ("KAPITEL 3", "GESPRÄCHSFÜHRUNG - DIE KERNKOMPETENZ", [
            "3.1 Grundhaltungen nach Carl Rogers",
            "3.2 Aktives Zuhören mit Jugendlichen",
            "3.3 Motivierende Gesprächsführung (MI)",
            "3.4 Lösungsorientierte Kurzberatung",
            "3.5 Gewaltfreie Kommunikation (GfK)",
            "3.6 Systemische Fragetechniken",
            "3.7 Erstgespräch mit Jugendlichen",
            "3.8 Schwierige Gesprächssituationen",
            "3.9 Gruppenarbeit mit Jugendlichen",
        ]),
        ("KAPITEL 4", "INTERVENTIONEN UND METHODEN", [
            "4.1 Kognitive Verhaltenstherapie - Adaptierte Techniken",
            "4.2 Emotionsregulationsstrategien",
            "4.3 DBT-Skills für die Beratung",
            "4.4 Narrative Therapie",
            "4.5 Kreative und erlebnisorientierte Methoden",
            "4.6 Psychoedukation für Jugendliche",
            "4.7 Achtsamkeit und Meditation",
            "4.8 Stärken- und Ressourcenorientierung",
            "4.9 Sicherheitsplanung bei Suizidalität",
        ]),
        ("KAPITEL 5", "ELTERNARBEIT UND SYSTEMISCHE ZUSAMMENARBEIT", [
            "5.1 Elterngespräche führen",
            "5.2 Systemische Zusammenarbeit",
            "5.3 Lehrerberatung",
        ]),
        ("KAPITEL 6", "KRISEN UND NOTFÄLLE", [
            "6.1 Krisenintervention - Allgemeines Modell",
            "6.2 Akute Suizidalität",
            "6.3 Akute Aggression / Gewalt",
            "6.4 Disclosure - Schwerwiegende Enthüllungen",
            "6.5 Schulkrise",
        ]),
        ("KAPITEL 7", "DIAGNOSTISCHE HILFSMITTEL", [
            "7.1 Screenings",
            "7.2 Verhaltensbeobachtung",
            "7.3 Genogramm und Netzwerkkarte",
        ]),
        ("KAPITEL 8", "SELBSTFÜRSORGE UND PROFESSIONELLE ENTWICKLUNG", [
            "8.1 Sekundäre Traumatisierung und Burnout",
            "8.2 Supervision und Intervision",
            "8.3 Ethische Dilemmata",
            "8.4 Weiterbildung und Literatur",
        ]),
        ("KAPITEL 9", "SPEZIALTHEMEN", [
            "9.1 Social Media, Gaming und digitale Welten",
            "9.2 Radikalisierung und Extremismus",
            "9.3 Hochbegabung und Underachievement",
            "9.4 Psychosomatik bei Jugendlichen",
            "9.5 Trauer bei Jugendlichen",
        ]),
        ("KAPITEL 10", "WERKZEUGKASTEN - ARBEITSBLÄTTER UND VORLAGEN", [
            "25 Arbeitsblätter und Vorlagen für die tägliche Praxis",
        ]),
    ]

    for chapter_num, chapter_title, subsections in toc_entries:
        # Chapter heading
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{chapter_num}: {chapter_title}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 51, 102)

        # Subsections
        for sub in subsections:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(sub)
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(80, 80, 80)

    # Page break after TOC
    doc.add_page_break()


def add_preface(doc):
    """Add preface/introduction."""
    p = doc.add_paragraph()
    run = p.add_run("VORWORT")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    # Suppress auto page break
    pPr = p._element.get_or_add_pPr()
    pageBreak = parse_xml(f'<w:pageBreakBefore {nsdecls("w")} w:val="0"/>')
    pPr.append(pageBreak)

    doc.add_paragraph()

    paras = [
        "Dieses Praxishandbuch wurde als umfassendes Referenzdokument für Psychologen am CDSE "
        "(Centre de développement scolaire et d'éducation) in Luxemburg entwickelt. Es richtet "
        "sich an Fachkräfte, die im pädagogisch-psychologischen Dienst im schulischen Kontext "
        "mit Jugendlichen zwischen 10 und 18 Jahren arbeiten, die sozio-emotionale Schwierigkeiten "
        "zeigen.",

        "Das Handbuch versteht sich als Nachschlagewerk für die gesamte berufliche Laufbahn. Es "
        "verbindet wissenschaftlich fundierte Erkenntnisse aus der Entwicklungspsychologie, klinischen "
        "Psychologie, Pädagogik und Sozialen Arbeit mit konkreten, alltagstauglichen Handlungsanleitungen.",

        "Ein zentrales Anliegen dieses Handbuchs ist die Rollenklarheit: Als Psychologe am CDSE führen "
        "Sie keine klinische Psychotherapie im engeren Sinne durch. Ihre Arbeit umfasst psychologische "
        "Beratung, Prävention, Förderung, Krisenintervention, Gruppenarbeit sowie Eltern- und "
        "Lehrerberatung. Dieses Handbuch respektiert diese Rollengrenze durchgehend und macht an den "
        "relevanten Stellen deutlich, wo die Grenzen der eigenen Kompetenz liegen und wann eine "
        "Überweisung an klinische Fachkräfte angezeigt ist.",

        "Die Inhalte berücksichtigen den besonderen luxemburgischen Kontext: die Multikulturalität und "
        "Mehrsprachigkeit des Landes, die spezifischen rechtlichen Rahmenbedingungen (luxemburgisches "
        "Recht, DSGVO, Kinderschutzgesetzgebung) sowie die lokalen institutionellen Strukturen und "
        "Anlaufstellen.",

        "Jedes Kapitel enthält neben theoretischen Grundlagen vor allem praktische Werkzeuge: "
        "Gesprächsleitfäden, Wort-für-Wort-Skripte, Checklisten, Entscheidungsbäume, Arbeitsblätter "
        "und Fallbeispiele. Das Ziel ist, dass Sie an jedem Arbeitstag in diesem Handbuch nachschlagen "
        "können und eine konkrete, anwendbare Antwort auf Ihre Frage finden.",

        "Nutzen Sie dieses Handbuch als lebendiges Dokument - ergänzen Sie es mit eigenen Erfahrungen, "
        "Notizen und lokalen Anpassungen. Die Arbeit mit Jugendlichen ist so vielfältig wie die "
        "Jugendlichen selbst, und kein Handbuch kann jede Situation abdecken. Aber es kann Ihnen ein "
        "solides Fundament bieten, auf dem Sie Ihre professionelle Praxis aufbauen."
    ]

    for text in paras:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)

    # Disclaimer
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("Hinweis zur Nutzung")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 51, 102)

    disclaimer_text = (
        "Dieses Handbuch ersetzt keine fachliche Supervision, Weiterbildung oder kollegiale Beratung. "
        "Die enthaltenen Informationen zu Störungsbildern dienen dem Erkennen und Verstehen, nicht der "
        "Diagnosestellung. Bei Unsicherheiten bezüglich klinischer Fragestellungen oder rechtlicher "
        "Aspekte konsultieren Sie bitte die zuständigen Fachstellen. Die beschriebenen Interventionen "
        "und Techniken sollen im Rahmen der eigenen Qualifikation und Rolle angewendet werden."
    )
    p2 = doc.add_paragraph(disclaimer_text)
    p2.paragraph_format.space_after = Pt(12)
    for run in p2.runs:
        run.font.italic = True

    doc.add_page_break()
