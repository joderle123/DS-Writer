"""
Styles and formatting utilities for the Praxishandbuch document.
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def setup_styles(doc):
    """Set up all custom styles for the document."""
    style = doc.styles

    # Normal style
    normal = style['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.15

    # Heading 1
    h1 = style['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 51, 102)
    h1.paragraph_format.space_before = Pt(24)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.page_break_before = True

    # Heading 2
    h2 = style['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(18)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 76, 153)
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(8)

    # Heading 3
    h3 = style['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(14)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 102, 153)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Heading 4
    h4 = style['Heading 4']
    h4.font.name = 'Calibri'
    h4.font.size = Pt(12)
    h4.font.bold = True
    h4.font.italic = True
    h4.font.color.rgb = RGBColor(0, 102, 153)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(4)

    # Create "Praxisbox" style for practice examples
    try:
        praxis_style = style.add_style('Praxisbox', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        praxis_style = style['Praxisbox']
    praxis_style.font.name = 'Calibri'
    praxis_style.font.size = Pt(10.5)
    praxis_style.font.italic = True
    praxis_style.font.color.rgb = RGBColor(51, 51, 51)
    praxis_style.paragraph_format.left_indent = Cm(1.5)
    praxis_style.paragraph_format.right_indent = Cm(1.0)
    praxis_style.paragraph_format.space_before = Pt(6)
    praxis_style.paragraph_format.space_after = Pt(6)

    # Create "Dialog" style
    try:
        dialog_style = style.add_style('Dialog', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        dialog_style = style['Dialog']
    dialog_style.font.name = 'Calibri'
    dialog_style.font.size = Pt(10.5)
    dialog_style.paragraph_format.left_indent = Cm(2.0)
    dialog_style.paragraph_format.space_before = Pt(2)
    dialog_style.paragraph_format.space_after = Pt(2)

    # Red Flag style
    try:
        redflag_style = style.add_style('RedFlag', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        redflag_style = style['RedFlag']
    redflag_style.font.name = 'Calibri'
    redflag_style.font.size = Pt(11)
    redflag_style.font.bold = True
    redflag_style.font.color.rgb = RGBColor(204, 0, 0)
    redflag_style.paragraph_format.left_indent = Cm(1.0)
    redflag_style.paragraph_format.space_before = Pt(6)
    redflag_style.paragraph_format.space_after = Pt(6)

    # Checklist style
    try:
        check_style = style.add_style('Checklist', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        check_style = style['Checklist']
    check_style.font.name = 'Calibri'
    check_style.font.size = Pt(10.5)
    check_style.paragraph_format.left_indent = Cm(1.5)
    check_style.paragraph_format.space_before = Pt(2)
    check_style.paragraph_format.space_after = Pt(2)

    return doc


def add_praxisbox(doc, title, content):
    """Add a practice example box with border."""
    # Add title
    p = doc.add_paragraph()
    p.style = doc.styles['Praxisbox']
    run = p.add_run(f"📋 PRAXISBEISPIEL: {title}")
    run.bold = True
    run.italic = False
    run.font.color.rgb = RGBColor(0, 102, 51)

    # Add border via XML
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:top w:val="single" w:sz="4" w:space="4" w:color="009966"/>'
        '  <w:left w:val="single" w:sz="4" w:space="8" w:color="009966"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="009966"/>'
        '  <w:right w:val="single" w:sz="4" w:space="8" w:color="009966"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # Add content lines
    for line in content.split('\n'):
        p2 = doc.add_paragraph()
        p2.style = doc.styles['Praxisbox']
        p2.add_run(line)
        # Add similar border
        pPr2 = p2._element.get_or_add_pPr()
        pBdr2 = parse_xml(
            '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '  <w:left w:val="single" w:sz="4" w:space="8" w:color="009966"/>'
            '  <w:right w:val="single" w:sz="4" w:space="8" w:color="009966"/>'
            '</w:pBdr>'
        )
        pPr2.append(pBdr2)


def add_dialog(doc, exchanges):
    """Add a dialog with speaker labels."""
    for speaker, text in exchanges:
        p = doc.add_paragraph()
        p.style = doc.styles['Dialog']
        run_speaker = p.add_run(f"{speaker}: ")
        run_speaker.bold = True
        if speaker.startswith("Psycholog") or speaker.startswith("Berater"):
            run_speaker.font.color.rgb = RGBColor(0, 76, 153)
        else:
            run_speaker.font.color.rgb = RGBColor(102, 51, 0)
        p.add_run(text)


def add_redflag(doc, text):
    """Add a red flag warning."""
    p = doc.add_paragraph()
    p.style = doc.styles['RedFlag']
    p.add_run(f"⚠ RED FLAG: {text}")


def add_checklist(doc, items):
    """Add a checklist."""
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Checklist']
        p.add_run(f"☐ {item}")


def add_checked_list(doc, items):
    """Add a checklist with checked items."""
    for item in items:
        p = doc.add_paragraph()
        p.style = doc.styles['Checklist']
        p.add_run(f"☑ {item}")


def add_flowchart(doc, steps):
    """Add a text-based flowchart."""
    p = doc.add_paragraph()
    run = p.add_run("ENTSCHEIDUNGSBAUM / FLOWCHART:")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.5)
        if isinstance(step, tuple):
            # Decision point
            question, yes_path, no_path = step
            run = p.add_run(f"❓ {question}")
            run.bold = True
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(2.5)
            p2.add_run(f"→ JA: {yes_path}")
            p3 = doc.add_paragraph()
            p3.paragraph_format.left_indent = Cm(2.5)
            p3.add_run(f"→ NEIN: {no_path}")
        else:
            if i > 0:
                arrow = doc.add_paragraph()
                arrow.paragraph_format.left_indent = Cm(2.0)
                arrow.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = arrow.add_run("↓")
                run.font.size = Pt(14)
            p.add_run(f"▶ {step}")


def add_table(doc, headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Blue background
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="003366"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9.5)
            # Alternate row shading
            if row_idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)

    return table


def add_bullet_list(doc, items, level=0):
    """Add bullet points."""
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(1.5 + level * 1.0)
        p.paragraph_format.space_after = Pt(3)


def add_numbered_list(doc, items):
    """Add numbered list."""
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f"{i}. {item}")
        p.paragraph_format.left_indent = Cm(1.5)
        p.paragraph_format.space_after = Pt(3)


def add_important_box(doc, text):
    """Add an important note box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    run = p.add_run(f"⚡ WICHTIG: {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(153, 0, 0)
    # Add border
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:top w:val="single" w:sz="6" w:space="4" w:color="CC0000"/>'
        '  <w:left w:val="single" w:sz="6" w:space="8" w:color="CC0000"/>'
        '  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="CC0000"/>'
        '  <w:right w:val="single" w:sz="6" w:space="8" w:color="CC0000"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_tip_box(doc, text):
    """Add a tip/hint box."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    run = p.add_run(f"💡 TIPP: {text}")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 51)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        '<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '  <w:top w:val="single" w:sz="4" w:space="4" w:color="009933"/>'
        '  <w:left w:val="single" w:sz="4" w:space="8" w:color="009933"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="4" w:color="009933"/>'
        '  <w:right w:val="single" w:sz="4" w:space="8" w:color="009933"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


def add_section_intro(doc, text):
    """Add introductory text for a section."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(51, 51, 51)
    p.paragraph_format.space_after = Pt(8)
