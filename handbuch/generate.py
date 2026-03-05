#!/usr/bin/env python3
"""
Main script to generate the complete Praxishandbuch document.
"""
import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.section import WD_ORIENT

from styles import setup_styles
from cover_and_toc import add_cover_page, add_table_of_contents, add_preface
from chapter1 import add_chapter1
from chapter2 import add_chapter2
from chapter3 import add_chapter3
from chapter4 import add_chapter4
from chapter5 import add_chapter5
from chapter6 import add_chapter6
from chapter7 import add_chapter7
from chapter8 import add_chapter8
from chapter9 import add_chapter9
from chapter10 import add_chapter10


def setup_document():
    """Create and configure the document."""
    doc = Document()

    # Set up page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    # Set up styles
    doc = setup_styles(doc)

    return doc


def add_footer(doc):
    """Add page numbers to footer."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = 1  # Center

        # Add page number field
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._element.append(fldChar1)

        run2 = paragraph.add_run()
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2._element.append(instrText)

        run3 = paragraph.add_run()
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3._element.append(fldChar2)

        # Add " / " separator
        run4 = paragraph.add_run(' / ')
        run4.font.size = Pt(9)

        # Add total pages field
        run5 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run5._element.append(fldChar3)

        run6 = paragraph.add_run()
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = ' NUMPAGES '
        run6._element.append(instrText2)

        run7 = paragraph.add_run()
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run7._element.append(fldChar4)


def main():
    print("=== Praxishandbuch Generator ===")
    print("Erstelle Dokument...")

    doc = setup_document()

    print("  Deckblatt und Inhaltsverzeichnis...")
    add_cover_page(doc)
    add_table_of_contents(doc)
    add_preface(doc)

    print("  Kapitel 1: Grundlagen der Adoleszenz...")
    add_chapter1(doc)

    print("  Kapitel 2: Störungsbilder und Auffälligkeiten...")
    add_chapter2(doc)

    print("  Kapitel 3: Gesprächsführung...")
    add_chapter3(doc)

    print("  Kapitel 4: Interventionen und Methoden...")
    add_chapter4(doc)

    print("  Kapitel 5: Elternarbeit und Systemische Zusammenarbeit...")
    add_chapter5(doc)

    print("  Kapitel 6: Krisen und Notfälle...")
    add_chapter6(doc)

    print("  Kapitel 7: Diagnostische Hilfsmittel...")
    add_chapter7(doc)

    print("  Kapitel 8: Selbstfürsorge und Professionelle Entwicklung...")
    add_chapter8(doc)

    print("  Kapitel 9: Spezialthemen...")
    add_chapter9(doc)

    print("  Kapitel 10: Werkzeugkasten...")
    add_chapter10(doc)

    # Add footer with page numbers
    print("  Seitennummerierung...")
    add_footer(doc)

    # Save document
    output_path = os.path.join(os.path.dirname(__file__), '..',
                               'Praxishandbuch_CDSE_Luxemburg.docx')
    output_path = os.path.abspath(output_path)

    print(f"\nSpeichere Dokument: {output_path}")
    doc.save(output_path)
    print(f"Dokument erfolgreich erstellt!")
    print(f"Dateigröße: {os.path.getsize(output_path) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
